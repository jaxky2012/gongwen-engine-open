#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
标准公文引擎 V14 - 继承基类

支持：红头公文、下行文、平行文
"""

from datetime import datetime
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from base_engine import BaseGongwenEngine


class StandardDocumentEngine(BaseGongwenEngine):
    """标准公文引擎（下行文/平行文）"""

    DEFAULT_COMPANY = "某某某某集团有限公司"

    def __init__(self, seal_logo_path=None):
        super().__init__()
        self.seal_logo_path = seal_logo_path

    def add_red_header(self, company_name=None):
        """添加红头标志（35mm锚点）"""
        company = company_name or self.DEFAULT_COMPANY
        text = f"{company}文件"

        # 35mm 空行占位
        spacer = self.doc.add_paragraph()
        spacer.paragraph_format.space_before = Pt(99)  # 35mm ≈ 99pt
        spacer.paragraph_format.space_after = Pt(0)
        spacer.paragraph_format.line_spacing = Pt(1)
        spacer.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY

        # 红头标志
        para = self.doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        para.paragraph_format.space_before = Pt(40)  # 防剃头
        para.paragraph_format.line_spacing = Pt(50)
        para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
        self._disable_grid_snap(para)

        run = para.add_run(text)
        self._set_font(run, name=self.PARAMS['font_title'], size=36, color_rgb=(255, 0, 0))

        # 字符压缩（超长名称）
        if len(text) > 14:
            scale = max(60, int(100 * 14 / len(text)))
            from docx.oxml import OxmlElement
            rPr = run._element.get_or_add_rPr()
            w_elem = OxmlElement('w:w')
            w_elem.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val', str(scale))
            rPr.append(w_elem)

    def add_red_separator_line(self):
        """添加红色分隔线"""
        self._add_horizontal_line(color_hex="FF0000", thickness="6", space_before_pt=11)

    def add_doc_number(self, doc_number):
        """添加发文字号（居中）"""
        if not doc_number:
            return
        para = self.doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        para.paragraph_format.space_after = Pt(0)
        run = para.add_run(doc_number)
        self._set_font(run)

    def add_title(self, title):
        """添加标题"""
        para = self.doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.add_run(title)
        self._set_font(run, name=self.PARAMS['font_title'], size=self.PARAMS['size_title'])

    def add_recipient(self, recipient):
        """添加主送机关"""
        if not recipient:
            return
        para = self.doc.add_paragraph()
        para.paragraph_format.first_line_indent = Pt(0)
        text = recipient.strip().rstrip(';:：；') + '：'
        run = para.add_run(text)
        self._set_font(run)

    def add_content(self, content_blocks):
        """添加正文"""
        for block in content_blocks:
            level = block.get('level', 0)
            text = block.get('text', '')
            index = block.get('index', '')

            para = self.doc.add_paragraph()

            if level == 0:
                para.paragraph_format.first_line_indent = self.PARAMS['indent_body']
                run = para.add_run(text)
                self._set_font(run)
            elif level == 1:
                content = f"{index}、{text}" if index else text
                run = para.add_run(content)
                self._set_font(run, name=self.PARAMS['font_l1'], bold=True)
            elif level == 2:
                para.paragraph_format.first_line_indent = self.PARAMS['indent_body']
                content = f"（{index}）{text}" if index else text
                run = para.add_run(content)
                self._set_font(run, name=self.PARAMS['font_l2'])
            elif level == 3:
                para.paragraph_format.first_line_indent = self.PARAMS['indent_body']
                content = f"{index}.{text}" if index else text
                run = para.add_run(content)
                self._set_font(run)
            elif level == 4:
                para.paragraph_format.first_line_indent = self.PARAMS['indent_body']
                content = f"({index}){text}" if index else text
                run = para.add_run(content)
                self._set_font(run)

    def add_signature(self, signatures, date_str=None):
        """添加落款"""
        date = date_str or datetime.now().strftime('%Y年%-m月%-d日')
        for sig in signatures:
            para = self.doc.add_paragraph()
            para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            para.paragraph_format.right_indent = Pt(self.PARAMS['size_body'] * 4)
            run = para.add_run(sig)
            self._set_font(run, name=self.PARAMS['font_sign'])

        date_para = self.doc.add_paragraph()
        date_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        date_para.paragraph_format.right_indent = Pt(self.PARAMS['size_body'] * 4.5)
        date_run = date_para.add_run(date)
        self._set_font(date_run, name=self.PARAMS['font_sign'])

    def add_banji(self, copy_to=None, issuer=None, issue_date=None):
        """添加版记"""
        self.doc.add_paragraph()

        # 版记首线
        self._add_horizontal_line(color_hex="000000", thickness="18", space_before_pt=0)

        if copy_to:
            para = self.doc.add_paragraph()
            para.paragraph_format.left_indent = Pt(14)
            run1 = para.add_run("抄送：")
            self._set_font(run1, size=self.PARAMS['size_banji'])
            run2 = para.add_run(copy_to)
            self._set_font(run2, size=self.PARAMS['size_banji'])

            # 内部分隔线
            self._add_horizontal_line(color_hex="000000", thickness="6", space_before_pt=0)

        issuer = issuer or self.DEFAULT_COMPANY + "办公室"
        issue_date = issue_date or datetime.now().strftime('%Y年%-m月%-d日')

        para = self.doc.add_paragraph()
        para.paragraph_format.left_indent = Pt(14)
        run1 = para.add_run(f" {issuer}")
        self._set_font(run1, size=self.PARAMS['size_banji'])
        para.add_run(" ")
        run2 = para.add_run(f"{issue_date}印发")
        self._set_font(run2, size=self.PARAMS['size_banji'])

        # 版记末线
        self._add_horizontal_line(color_hex="000000", thickness="18", space_before_pt=0)


def generate_gongwen(data, output_path=None, seal_path=None):
    """生成公文（便捷接口）"""
    import os
    script_dir = os.path.dirname(os.path.abspath(__file__))
    default_seal = os.path.join(script_dir, "2.jpg")
    use_seal = seal_path or data.get('seal_path') or default_seal

    engine = StandardDocumentEngine(seal_logo_path=use_seal)

    engine.add_red_header(data.get('company_name'))
    engine.doc.add_paragraph()
    engine.doc.add_paragraph()
    engine.add_doc_number(data.get('doc_number'))
    engine.add_red_separator_line()
    engine.doc.add_paragraph()
    engine.doc.add_paragraph()

    if data.get('title'):
        engine.add_title(data['title'])

    engine.doc.add_paragraph()
    engine.add_recipient(data.get('recipient'))

    if data.get('content_blocks'):
        engine.add_content(data['content_blocks'])

    signatures = data.get('signatures', [data.get('company_name', engine.DEFAULT_COMPANY)])
    engine.add_signature(signatures, data.get('date'))
    engine.add_banji(copy_to=data.get('copy_to'), issuer=data.get('issuer'), issue_date=data.get('issue_date'))
    engine.add_page_number()

    if not output_path:
        title_safe = data.get('title', '公文')[:20].replace('/', '_')
        output_path = f"公文_{title_safe}_{datetime.now().strftime('%Y%m%d')}.docx"

    return engine.save(output_path)


if __name__ == "__main__":
    test_data = {
        "company_name": "某某某某集团有限公司",
        "doc_number": "某企〔2026〕1号",
        "title": "关于开展2026年度工作的通知",
        "recipient": "所属各单位",
        "content_blocks": [
            {"level": 0, "text": "为进一步推动工作落实，现将有关事项通知如下。"},
            {"level": 1, "index": "一", "text": "工作目标"},
            {"level": 0, "text": "全年实现高质量发展目标。"},
        ],
        "signatures": ["某某某某集团有限公司"],
        "date": "2026年4月20日"
    }

    result = generate_gongwen(test_data)
    print(f"✅ 公文已生成: {result}")
