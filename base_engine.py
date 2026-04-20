#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
公文渲染基类引擎 - 所有公文类型的公共底座

职责：
- 页面设置（GB/T 9704-2012：37/35/28/26mm 边距）
- 字体渲染（强制中文字体）
- 画横线（防止塌陷的核心修复）
- 页码生成
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


class BaseGongwenEngine:
    """
    公文渲染基类

    物理参数体系（GB/T 9704-2012）：
    - 版心宽度：156mm
    - 行距：28pt（22行/页）
    - 天头：37mm
    - 下白边：35mm
    - 订口：28mm
    - 切口：26mm
    """

    # 公共参数体系
    PARAMS = {
        'font_title': 'STZhongsong',  # 小标宋
        'font_body': 'FangSong',  # 仿宋
        'font_l1': 'SimHei',  # 一级标题：黑体
        'font_l2': 'STKaiti',  # 二级标题：楷体
        'font_sign': 'STKaiti',  # 签名：楷体
        'size_title': 22,  # 二号
        'size_body': 16,  # 三号
        'size_banji': 14,  # 四号
        'line_spacing': 28,  # 固定行距
        'content_width': Cm(15.6),  # 版心宽度
        'indent_body': Pt(32),  # 正文首行缩进
    }

    # 页面边距（GB/T 9704-2012）
    PAGE_MARGIN = {
        'top': Cm(3.7),
        'bottom': Cm(3.5),
        'left': Cm(2.8),
        'right': Cm(2.6),
    }

    def __init__(self):
        """初始化基类"""
        self.doc = Document()
        self._setup_page()

    def _setup_page(self):
        """设置页面参数"""
        sec = self.doc.sections[0]
        sec.top_margin = self.PAGE_MARGIN['top']
        sec.bottom_margin = self.PAGE_MARGIN['bottom']
        sec.left_margin = self.PAGE_MARGIN['left']
        sec.right_margin = self.PAGE_MARGIN['right']

    def _set_font(self, run, name=None, size=None, color_rgb=None, bold=False):
        """
        强制中文字体渲染

        Args:
            run: docx Run 对象
            name: 字体名称
            size: 字号（磅）
            color_rgb: 颜色元组 (R, G, B)
            bold: 是否加粗
        """
        f_name = name or self.PARAMS['font_body']
        f_size = size or self.PARAMS['size_body']

        run.font.name = f_name
        run._element.rPr.rFonts.set(qn('w:eastAsia'), f_name)
        run.font.size = Pt(f_size)
        run.bold = bold

        if color_rgb:
            run.font.color.rgb = RGBColor(*color_rgb)

    def _add_horizontal_line(self, color_hex="000000", thickness="6", space_before_pt=0):
        """
        [核心修复] 添加水平分隔线

        解决 Word/WPS 中空段落边框塌陷的问题：
        1. 注入不可见空格（半角空格）撑开段落
        2. 字号设为最小（1pt），不影响视觉高度
        3. sz 参数精确控制线宽（单位：1/8 磅）

        Args:
            color_hex: 十六进制颜色，如 "FF0000"（红色）、"00AA00"（绿色）、"000000"（黑色）
            thickness: 线宽参数，"6"=0.75磅（细线），"12"=1.5磅（粗线）
            space_before_pt: 段前间距（磅）
        """
        para = self.doc.add_paragraph()
        para.paragraph_format.space_before = Pt(space_before_pt)
        para.paragraph_format.space_after = Pt(0)
        para.paragraph_format.line_spacing = Pt(1)
        para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY

        # [关键] 注入不可见字符，防止边框塌陷
        run = para.add_run(" ")
        run.font.size = Pt(1)  # 最小字号，不影响排版

        # 注入底层边框 XML
        pPr = para._p.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), thickness)
        bottom.set(qn('w:color'), color_hex)
        bottom.set(qn('w:space'), '0')  # 消除边框与段落间距
        pBdr.append(bottom)
        pPr.append(pBdr)

    def _disable_grid_snap(self, paragraph):
        """禁用文档网格吸附"""
        pPr = paragraph._p.get_or_add_pPr()
        snap = OxmlElement('w:snapToGrid')
        snap.set(qn('w:val'), '0')
        pPr.append(snap)

    def add_page_number(self):
        """添加页码（四号宋体，居中）"""
        for section in self.doc.sections:
            footer = section.footer
            footer.is_linked_to_previous = False
            para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # 左侧短横线
            run1 = para.add_run("— ")
            self._set_font(run1, size=self.PARAMS['size_banji'])

            # PAGE 域代码
            fldChar1 = OxmlElement('w:fldChar')
            fldChar1.set(qn('w:fldCharType'), 'begin')
            instrText = OxmlElement('w:instrText')
            instrText.text = "PAGE"
            fldChar2 = OxmlElement('w:fldChar')
            fldChar2.set(qn('w:fldCharType'), 'end')

            run2 = para.add_run()
            run2._r.append(fldChar1)
            run2._r.append(instrText)
            run2._r.append(fldChar2)
            self._set_font(run2, size=self.PARAMS['size_banji'])

            # 右侧短横线
            run3 = para.add_run(" —")
            self._set_font(run3, size=self.PARAMS['size_banji'])

    def save(self, output_path):
        """保存文档"""
        self.doc.save(output_path)
        return output_path
