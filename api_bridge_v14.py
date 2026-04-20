# -*- coding: utf-8 -*-
# ==============================================================================
#  GongWen-V14 Pixel-Perfect Render Engine
#  Architecture: PUA (Plan-Ultimate-Action) Multi-Agent Framework
#  Developer: Jeff M. (AI Productivity Coach)
# ==============================================================================

import os
from datetime import datetime
from fastapi import FastAPI
from fastapi.responses import FileResponse
from pydantic import BaseModel
from typing import List, Optional

from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT

# ------------------------------------------------------------------------------
# 1. FastAPI 接口名片定义 (脱敏版)
# ------------------------------------------------------------------------------
app = FastAPI(
    title="GongWen-V14 智能公文排版引擎",
    description="基于 PUA 架构的像素级公文物理渲染中枢。严格遵循 661号公文排版规范。",
    version="14.0.0",
    contact={
        "name": "Jeff M. (AI Productivity Coach)",
        "email": "jeffm8888@gmail.com",
    }
)

# ------------------------------------------------------------------------------
# 2. Pydantic 数据载荷 Schema (约束大模型输出)
# ------------------------------------------------------------------------------
class ContentBlock(BaseModel):
    level: int
    index: Optional[str] = None
    text: str
    children: Optional[List['ContentBlock']] = []

class DocRequest(BaseModel):
    doc_type: str = "普通发文"
    issue_agencies: List[str]  # 发文机关列表 (红头)
    doc_number: str            # 发文字号 (仅限第一主办机关)
    issuer: Optional[str] = None # 签发人 (上行文用)
    title: str                 # 公文标题
    recipients: str            # 主送机关
    content_blocks: List[ContentBlock] # 正文层级树
    signatures: List[str]      # 落款机关
    issue_date: str            # 成文日期
    cc_orgs: str               # 抄送机关
    print_agency: str          # 印发机关

# ------------------------------------------------------------------------------
# 3. V14 引擎底层 XML 操控组件
# ------------------------------------------------------------------------------
def add_official_run(paragraph, text, font_family='仿宋_GB2312', size_pt=16, color_rgb=(0, 0, 0), bold=False):
    """双轨字体挂载与色彩锁定"""
    run = paragraph.add_run(text)
    run.font.name = font_family
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_family)
    run.font.size = Pt(size_pt)
    run.font.color.rgb = RGBColor(*color_rgb)
    run.font.bold = bold
    return run

def fix_tito_issue(paragraph, line_spacing_pt=28):
    """解决大字号'剃头'现象的物理对齐"""
    pf = paragraph.paragraph_format
    pf.space_before = Pt(3) # 注入段前间距补偿
    pf.line_spacing = Pt(line_spacing_pt)
    pf.line_spacing_rule = WD_LINE_SPACING.EXACT
    
    # 强制脱离文档网格
    pPr = paragraph._p.get_or_add_pPr()
    snap = OxmlElement('w:snapToGrid')
    snap.set(qn('w:val'), '0')
    pPr.append(snap)

def inject_pBdr_xml(paragraph, border_type='top', sz=6):
    """版记物理隔离线注入 (sz=6为0.75pt, sz=2为0.25pt)"""
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = pPr.find(qn('w:pBdr'))
    if pBdr is None:
        pBdr = OxmlElement('w:pBdr')
        pPr.append(pBdr)
    
    border = OxmlElement(f'w:{border_type}')
    border.set(qn('w:val'), 'single')
    border.set(qn('w:sz'), str(sz))
    border.set(qn('w:space'), '1')
    border.set(qn('w:color'), 'FF0000' if border_type == 'bottom' and sz == 6 else '000000') # 红头线为红色
    pBdr.append(border)

def apply_silent_branding(doc):
    """V14 引擎专属：物理层核心属性注入 (Jeff M. 专属版权)"""
    props = doc.core_properties
    props.author = "Jeff M. (AI Productivity Coach)"
    props.comments = (
        "本文件由智能公文排版引擎 V14 像素级渲染生成。\n"
        "架构师：Jeff M.\n"
        "Email: jeffm8888@gmail.com | QQ: 93710348\n"
        "技术架构：基于 PUA 框架的多智能体协同渲染。"
    )
    props.title = "GongWen-V14 智能排版系统输出"
    props.category = "Open Source Tooling"
    props.content_status = "V14.0.0-Stable"
    props.last_modified_by = "Jeff M."

# ------------------------------------------------------------------------------
# 4. 核心渲染流程
# ------------------------------------------------------------------------------
def render_document(data: DocRequest, output_path: str):
    doc = Document()
    
    # [页面设置] 预留给未来添加确切的页边距代码 (上37mm等)
    # section = doc.sections[0]
    
    # --- 1. 发文机关标志 (红头) ---
    agency_count = len(data.issue_agencies)
    font_sizes = {1: Pt(48), 2: Pt(36), 3: Pt(26), 4: Pt(22), 5: Pt(18)}
    current_size = font_sizes.get(agency_count, Pt(16))
    
    for agency in data.issue_agencies:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        fix_tito_issue(p, line_spacing_pt=35)
        add_official_run(p, agency, font_family='方正小标宋简体', size_pt=current_size.pt, color_rgb=(255, 0, 0))
    
    p_file = doc.add_paragraph()
    p_file.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_official_run(p_file, "文件", font_family='方正小标宋简体', size_pt=current_size.pt, color_rgb=(255, 0, 0))
    
    # 注入红线
    inject_pBdr_xml(p_file, border_type='bottom', sz=6)
    
    # --- 2. 发文字号 ---
    p_num = doc.add_paragraph()
    p_num.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_num.paragraph_format.space_before = Pt(10)
    add_official_run(p_num, data.doc_number, font_family='仿宋_GB2312', size_pt=16)
    
    # --- 3. 公文标题 ---
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_title.paragraph_format.space_before = Pt(24)
    p_title.paragraph_format.space_after = Pt(24)
    add_official_run(p_title, data.title, font_family='方正小标宋简体', size_pt=22)
    
    # --- 4. 主送机关 ---
    p_rec = doc.add_paragraph()
    add_official_run(p_rec, f"{data.recipients}：", font_family='仿宋_GB2312', size_pt=16)
    
    # --- 5. 正文渲染 (递归) ---
    def render_blocks(blocks: List[ContentBlock]):
        for block in blocks:
            p = doc.add_paragraph()
            p.paragraph_format.first_line_indent = Pt(32) # 首行缩进2字符
            p.paragraph_format.line_spacing = Pt(28)
            p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACT
            
            text_content = f"{block.index}、{block.text}" if block.level == 1 else block.text
            
            if block.level == 1:
                add_official_run(p, text_content, font_family='黑体', size_pt=16)
            elif block.level == 2:
                add_official_run(p, text_content, font_family='楷体_GB2312', size_pt=16)
            else:
                add_official_run(p, text_content, font_family='仿宋_GB2312', size_pt=16)
            
            if block.children:
                render_blocks(block.children)
                
    render_blocks(data.content_blocks)
    
    # --- 6. 落款与日期 ---
    p_date = doc.add_paragraph()
    p_date.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    add_official_run(p_date, f"{data.issue_date}    ", font_family='仿宋_GB2312', size_pt=16) # 右空四字
    
    # 注入 Jeff M. 底层版权属性
    apply_silent_branding(doc)
    
    doc.save(output_path)
    return output_path

# ------------------------------------------------------------------------------
# 5. FastAPI 路由
# ------------------------------------------------------------------------------
@app.post("/api/gongwen/v14/generate")
async def generate_gongwen(request: DocRequest):
    # 确保输出目录存在
    output_dir = os.path.join(os.getcwd(), "output")
    os.makedirs(output_dir, exist_ok=True)
    
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    filename = f"V14_{request.doc_type}_{timestamp}.docx"
    filepath = os.path.join(output_dir, filename)
    
    # 调用 V14 引擎渲染
    render_document(request, filepath)
    
    # 返回生成的文件
    return FileResponse(
        path=filepath, 
        filename=filename, 
        media_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )

if __name__ == "__main__":
    import uvicorn
    # 启动服务，监听所有 IP 的 8000 端口
    uvicorn.run(app, host="0.0.0.0", port=8000)